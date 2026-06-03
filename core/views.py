"""
View-функции приложения core.
Требования: F01–F14 (1.4_Analiz_Trebovaniy_Polzovateley.md)
"""

from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import logout as auth_logout, login as auth_login
from django.contrib.auth.decorators import login_required
from django.views.decorators.http import require_POST
from django.contrib import messages
from django.urls import reverse

from django.core.paginator import Paginator

import os
from django.http import FileResponse, Http404, HttpResponse

from .models import Service, News, Project, Role, Notification, Attachment
from .forms import (
    ContactMessageForm, LoginForm, RegisterForm,
    ProjectCreateForm, ProfileUpdateForm,
    TicketFilterForm, TicketStatusForm, CommentForm, AttachmentUploadForm,
)

# ─── Константы ──────────────────────────────────────────────────────────────

# Роль по умолчанию для самостоятельной регистрации (F01)
ROLE_CLIENT = 'client'

# Количество записей на главной странице (ВКР-033)
HOME_SERVICES_COUNT = 3
HOME_NEWS_COUNT     = 3

# Сообщение об успешной отправке обращения (ВКР-035)
CONTACT_SUCCESS_MESSAGE = (
    'Ваше обращение принято. Мы свяжемся с вами в ближайшее время.'
)

# Сообщения аутентификации (ВКР-036)
REGISTER_SUCCESS_MESSAGE = 'Добро пожаловать! Аккаунт успешно создан.'
REGISTER_ROLE_NOT_FOUND  = (
    'Не удалось назначить роль. Обратитесь к администратору.'
)

# Лимиты страниц контента (ВКР-034)
SERVICES_PER_PAGE  = 20
NEWS_PER_PAGE      = 10
PROJECTS_PER_PAGE  = 9   # кратно 3 — сетка 3-в-ряд

# Дашборд клиента (ВКР-038)
DASHBOARD_TICKETS_PER_PAGE = 10
TICKET_CREATE_SUCCESS_MESSAGE = 'Заявка успешно создана. Мы свяжемся с вами в ближайшее время.'

# Детальная заявка (ВКР-040)
TICKET_NOT_FOUND_MESSAGE = 'Заявка не найдена или недоступна.'

# Профиль и уведомления (ВКР-041)
PROFILE_UPDATE_SUCCESS  = 'Профиль успешно обновлён.'
NOTIFICATIONS_PER_PAGE  = 20

# ЛК менеджера (ВКР-043/044)
MANAGER_TICKETS_PER_PAGE  = 15
STATUS_UPDATED_MESSAGE    = 'Статус заявки обновлён.'
COMMENT_ADDED_MESSAGE     = 'Комментарий добавлен.'

# Файловая система (ВКР-045)
ATTACHMENT_UPLOAD_SUCCESS = 'Файл успешно прикреплён к заявке.'


# ─── Публичные страницы ──────────────────────────────────────────────────────

def home_view(request):
    """
    Главная страница — доступна всем.
    Передаёт активные услуги и последние новости (ВКР-033).
    Требования: F02 (точка входа для подачи заявки через контакты).
    """
    services = (
        Service.objects
        .filter(is_active=True)
        .order_by('title')[:HOME_SERVICES_COUNT]
    )
    news = (
        News.objects
        .filter(published_at__isnull=False)
        .order_by('-published_at')
        .select_related('author')[:HOME_NEWS_COUNT]
    )
    return render(request, 'core/home.html', {
        'services': services,
        'news':     news,
    })


def services_view(request):
    """
    Каталог услуг — полный список активных услуг, сгруппированных по категории.
    Доступна всем. Требования: F02 (точка входа к услугам перед подачей заявки).
    """
    services = (
        Service.objects
        .filter(is_active=True)
        .select_related('category')
        .order_by('category__name', 'title')[:SERVICES_PER_PAGE]
    )
    return render(request, 'core/services.html', {'services': services})


def news_view(request):
    """
    Лента новостей компании — только опубликованные записи.
    Доступна всем.
    """
    news_items = (
        News.objects
        .filter(published_at__isnull=False)
        .select_related('author')
        .order_by('-published_at')[:NEWS_PER_PAGE]
    )
    return render(request, 'core/news.html', {'news_items': news_items})


def projects_view(request):
    """
    Публичное портфолио — завершённые и закрытые заявки.
    Личные данные клиента не раскрываются (ВКР-034).
    """
    projects = (
        Project.objects
        .filter(status__in=Project.PORTFOLIO_STATUSES)
        .select_related('service', 'service__category')
        .order_by('-created_at')[:PROJECTS_PER_PAGE]
    )
    return render(request, 'core/projects.html', {'projects': projects})


def contacts_view(request):
    """
    Страница контактов с формой обратной связи (ВКР-035).
    GET  — отображает форму (поля name/email предзаполнены для авторизованных).
    POST — сохраняет ContactMessage, redirect-after-POST → success banner.
    Доступна всем ролям и анонимным пользователям (F02).
    """
    if request.method == 'POST':
        form = ContactMessageForm(request.POST)
        if form.is_valid():
            contact = form.save(commit=False)
            if request.user.is_authenticated:
                contact.user = request.user
            contact.save()
            messages.success(request, CONTACT_SUCCESS_MESSAGE)
            return redirect(reverse('core:contacts'))
        # Невалидная форма — повторный рендер с ошибками (без редиректа)
        return render(request, 'core/contacts.html', {'form': form})

    # GET — предзаполнение для авторизованных пользователей
    initial = {}
    if request.user.is_authenticated:
        first = request.user.first_name.strip()
        last  = request.user.last_name.strip()
        initial['name']  = f'{first} {last}'.strip() or request.user.username
        initial['email'] = request.user.email

    form = ContactMessageForm(initial=initial)
    return render(request, 'core/contacts.html', {'form': form})


# ─── Статические публичные страницы (ВКР-037) ───────────────────────────────

def about_view(request):
    """
    О компании — статическая публичная страница.
    Доступна всем. Закрывает критерий §4 «≥10 страниц».
    """
    return render(request, 'core/about.html')


def sitemap_view(request):
    """
    Карта сайта — структурированный список всех разделов.
    Доступна всем. Закрывает критерий §4 «≥10 страниц».
    """
    return render(request, 'core/sitemap.html')


def news_detail_view(request, pk):
    """
    Детальная страница новости.
    Доступна всем. 404 если новость не опубликована (published_at is None).
    """
    news_item = get_object_or_404(
        News.objects.select_related('author'),
        pk=pk,
        published_at__isnull=False,
    )
    return render(request, 'core/news_detail.html', {'news_item': news_item})


# ─── Аутентификация (F01) ────────────────────────────────────────────────────

def login_view(request):
    """
    Форма входа (F01). Только для анонимных пользователей.
    POST: authenticate() → login() → redirect to ?next= или LOGIN_REDIRECT_URL.
    """
    if request.user.is_authenticated:
        return redirect('core:home')

    if request.method == 'POST':
        form = LoginForm(request, data=request.POST)
        if form.is_valid():
            user = form.get_user()
            auth_login(request, user)
            next_url = request.POST.get('next') or request.GET.get('next') or 'core:home'
            return redirect(next_url)
    else:
        form = LoginForm(request)

    return render(request, 'core/login.html', {
        'form': form,
        'next': request.GET.get('next', ''),
    })


def register_view(request):
    """
    Форма регистрации (F01). Только для анонимных пользователей.
    Назначает роль ROLE_CLIENT, хэширует пароль, автоматически выполняет вход.
    """
    if request.user.is_authenticated:
        return redirect('core:home')

    if request.method == 'POST':
        form = RegisterForm(request.POST)
        if form.is_valid():
            try:
                role_client = Role.objects.get(name=ROLE_CLIENT)
            except Role.DoesNotExist:
                messages.error(request, REGISTER_ROLE_NOT_FOUND)
                return render(request, 'core/register.html', {'form': form})

            user = form.save(commit=False)
            user.role = role_client
            user.save()
            auth_login(request, user)
            messages.success(request, REGISTER_SUCCESS_MESSAGE)
            return redirect('core:home')
    else:
        form = RegisterForm()

    return render(request, 'core/register.html', {'form': form})


@require_POST
def logout_view(request):
    """
    Выход из системы — только POST (защита от CSRF-logout через GET).
    Редирект на LOGIN_URL определён в settings.LOGOUT_REDIRECT_URL.
    """
    auth_logout(request)
    return redirect('core:login')


# ─── Личный кабинет (F02–F10) ───────────────────────────────────────────────

@login_required
def dashboard_view(request):
    """
    Дашборд клиента — список его заявок с пагинацией и счётчиками по статусам (F03).
    Требует авторизации. Неавторизованных редиректит на /login/?next=/dashboard/.
    Полная реализация: ВКР-038.
    """
    tickets_qs = (
        Project.objects
        .filter(user=request.user)
        .select_related('service', 'priority')
        .order_by('-created_at')
    )

    # Счётчики по статусам для stat-cards
    stats = {
        'new':         tickets_qs.filter(status=Project.STATUS_NEW).count(),
        'in_progress': tickets_qs.filter(status=Project.STATUS_IN_PROGRESS).count(),
        'resolved':    tickets_qs.filter(
                           status__in=[Project.STATUS_RESOLVED, Project.STATUS_CLOSED]
                       ).count(),
        'total':       tickets_qs.count(),
    }

    paginator = Paginator(tickets_qs, DASHBOARD_TICKETS_PER_PAGE)
    page_obj  = paginator.get_page(request.GET.get('page'))

    return render(request, 'core/dashboard.html', {
        'page_obj': page_obj,
        'stats':    stats,
    })


@login_required
def ticket_create_view(request):
    """
    Форма создания новой заявки клиентом (F02).
    POST: сохраняет заявку со статусом STATUS_NEW и текущим пользователем.
    Redirect-after-POST → дашборд.
    """
    if request.method == 'POST':
        form = ProjectCreateForm(request.POST)
        if form.is_valid():
            ticket = form.save(commit=False)
            ticket.user   = request.user
            ticket.status = Project.STATUS_NEW
            ticket.save()
            messages.success(request, TICKET_CREATE_SUCCESS_MESSAGE)
            return redirect(reverse('core:dashboard'))
        return render(request, 'core/ticket_create.html', {'form': form})

    form = ProjectCreateForm()
    return render(request, 'core/ticket_create.html', {'form': form})


# ─── Детальная заявка клиента (ВКР-040) ─────────────────────────────────────

@login_required
def ticket_detail_view(request, pk):
    """
    Детальная страница заявки клиента (F03, F05).
    Клиент видит только свои заявки — фильтр по user=request.user.
    Показывает: статус, описание, комментарии, список вложений (имена).
    """
    ticket = get_object_or_404(
        Project.objects
        .select_related('service', 'priority', 'manager')
        .prefetch_related('comments__user', 'attachments'),
        pk=pk,
        user=request.user,
    )
    return render(request, 'core/ticket_detail.html', {'ticket': ticket})


# ─── Профиль и уведомления (ВКР-041) ────────────────────────────────────────

@login_required
def profile_view(request):
    """
    Профиль пользователя — просмотр и редактирование (F01).
    GET: отображает форму с текущими данными.
    POST: сохраняет first_name, last_name, email.
    """
    if request.method == 'POST':
        form = ProfileUpdateForm(request.POST, instance=request.user)
        if form.is_valid():
            form.save()
            messages.success(request, PROFILE_UPDATE_SUCCESS)
            return redirect(reverse('core:profile'))
        return render(request, 'core/profile.html', {'form': form})

    form = ProfileUpdateForm(instance=request.user)
    return render(request, 'core/profile.html', {'form': form})


@login_required
def notifications_view(request):
    """
    Список уведомлений пользователя с пагинацией (F04).
    Показывает непрочитанные первыми.
    """
    notifications_qs = (
        Notification.objects
        .filter(user=request.user)
        .select_related('project')
        .order_by('is_read', '-created_at')
    )
    unread_count = notifications_qs.filter(is_read=False).count()

    paginator = Paginator(notifications_qs, NOTIFICATIONS_PER_PAGE)
    page_obj  = paginator.get_page(request.GET.get('page'))

    return render(request, 'core/notifications.html', {
        'page_obj':     page_obj,
        'unread_count': unread_count,
    })


@login_required
@require_POST
def notification_mark_read_view(request, pk):
    """
    Отметить уведомление прочитанным (POST).
    Проверяет владельца — чужое уведомление → 404.
    """
    notification = get_object_or_404(Notification, pk=pk, user=request.user)
    notification.is_read = True
    notification.save(update_fields=['is_read'])
    return redirect(reverse('core:notifications'))


# ─── ЛК Менеджера (ВКР-043/044) ─────────────────────────────────────────────

def manager_required(view_func):
    """Декоратор: требует роль manager или admin. Иначе → 403."""
    from functools import wraps
    from django.http import HttpResponseForbidden

    @wraps(view_func)
    def wrapper(request, *args, **kwargs):
        if not request.user.is_authenticated:
            return redirect(f'/login/?next={request.path}')
        if not (request.user.is_manager or request.user.is_admin):
            return HttpResponseForbidden('Доступ запрещён.')
        return view_func(request, *args, **kwargs)
    return wrapper


@manager_required
def manager_dashboard_view(request):
    """
    Дашборд менеджера — список всех заявок с фильтром по статусу и поиском (F06).
    """
    form = TicketFilterForm(request.GET or None)
    tickets_qs = (
        Project.objects
        .select_related('user', 'service', 'priority', 'manager')
        .order_by('-created_at')
    )

    if form.is_valid():
        status = form.cleaned_data.get('status')
        search = form.cleaned_data.get('search', '').strip()
        if status:
            tickets_qs = tickets_qs.filter(status=status)
        if search:
            tickets_qs = tickets_qs.filter(title__icontains=search)

    # Счётчики по статусам
    stats = {
        'new':         tickets_qs.filter(status=Project.STATUS_NEW).count(),
        'in_progress': tickets_qs.filter(status=Project.STATUS_IN_PROGRESS).count(),
        'resolved':    tickets_qs.filter(
                           status__in=[Project.STATUS_RESOLVED, Project.STATUS_CLOSED]
                       ).count(),
        'total':       tickets_qs.count(),
    }

    paginator = Paginator(tickets_qs, MANAGER_TICKETS_PER_PAGE)
    page_obj  = paginator.get_page(request.GET.get('page'))

    return render(request, 'core/manager_dashboard.html', {
        'page_obj': page_obj,
        'stats':    stats,
        'form':     form,
    })


@manager_required
def manager_ticket_detail_view(request, pk):
    """
    Детальная заявка менеджера (F07): смена статуса + добавление комментария.
    POST с action='status' — меняет статус.
    POST с action='comment' — добавляет комментарий.
    """
    ticket = get_object_or_404(
        Project.objects
        .select_related('user', 'service', 'priority', 'manager')
        .prefetch_related('comments__user', 'attachments'),
        pk=pk,
    )

    status_form  = TicketStatusForm(instance=ticket)
    comment_form = CommentForm()

    if request.method == 'POST':
        action = request.POST.get('action')

        if action == 'status':
            status_form = TicketStatusForm(request.POST, instance=ticket)
            if status_form.is_valid():
                status_form.save()
                # Назначаем менеджера если ещё не назначен
                if not ticket.manager:
                    ticket.manager = request.user
                    ticket.save(update_fields=['manager'])
                messages.success(request, STATUS_UPDATED_MESSAGE)
                return redirect(reverse('core:manager_ticket_detail', kwargs={'pk': pk}))

        elif action == 'comment':
            comment_form = CommentForm(request.POST)
            if comment_form.is_valid():
                comment = comment_form.save(commit=False)
                comment.project = ticket
                comment.user    = request.user
                comment.save()
                messages.success(request, COMMENT_ADDED_MESSAGE)
                return redirect(reverse('core:manager_ticket_detail', kwargs={'pk': pk}))

    return render(request, 'core/manager_ticket_detail.html', {
        'ticket':       ticket,
        'status_form':  status_form,
        'comment_form': comment_form,
        'upload_form':  AttachmentUploadForm(),
    })


@manager_required
def manager_stats_view(request):
    """
    Статистика по заявкам для менеджера (F13).
    """
    from django.db.models import Count

    status_counts = (
        Project.objects
        .values('status')
        .annotate(count=Count('id'))
        .order_by('status')
    )
    manager_counts = (
        Project.objects
        .filter(manager__isnull=False)
        .values('manager__username', 'manager__first_name', 'manager__last_name')
        .annotate(count=Count('id'))
        .order_by('-count')
    )

    return render(request, 'core/manager_stats.html', {
        'status_counts':  status_counts,
        'manager_counts': manager_counts,
        'status_labels':  dict(Project.STATUS_CHOICES),
    })


# ─── Файловая система (ВКР-045) ───────────────────────────────────────────────

@login_required
@require_POST
def attachment_upload_view(request, pk):
    """
    Загрузка файла к заявке (F02/F07).
    Клиент — только к своей заявке. Менеджер/admin — к любой.
    """
    if request.user.is_manager or request.user.is_admin:
        ticket = get_object_or_404(Project, pk=pk)
    else:
        ticket = get_object_or_404(Project, pk=pk, user=request.user)

    form = AttachmentUploadForm(request.POST, request.FILES)
    if form.is_valid():
        uploaded = form.cleaned_data['file']
        attachment = Attachment(
            project   = ticket,
            filename  = uploaded.name,
            file      = uploaded,
            file_type = uploaded.content_type or '',
            file_size = uploaded.size,
        )
        attachment.save()
        messages.success(request, ATTACHMENT_UPLOAD_SUCCESS)

    # Редиректим обратно на страницу заявки (клиент или менеджер)
    if request.user.is_manager or request.user.is_admin:
        return redirect(reverse('core:manager_ticket_detail', kwargs={'pk': pk}))
    return redirect(reverse('core:ticket_detail', kwargs={'pk': pk}))


@login_required
def attachment_download_view(request, pk):
    """
    Скачивание файла вложения (F02/F07).
    Клиент — только вложения своих заявок.
    Менеджер/admin — любые вложения.
    """
    attachment = get_object_or_404(Attachment.objects.select_related('project__user'), pk=pk)

    # Проверка доступа
    if not (request.user.is_manager or request.user.is_admin):
        if attachment.project.user != request.user:
            raise Http404

    if not attachment.file:
        raise Http404

    response = FileResponse(attachment.file.open('rb'), as_attachment=True)
    response['Content-Disposition'] = f'attachment; filename="{attachment.filename}"'
    return response


# ─── Экспорт отчётов (ВКР-046) ──────────────────────────────────────────────

@manager_required
def export_tickets_xlsx_view(request):
    """
    Экспорт всех заявок в .xlsx (F10).
    Доступно менеджеру и администратору.
    """
    import io
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment

    STATUS_LABELS = dict(Project.STATUS_CHOICES)
    HEADER_COLOR = 'FF2563EB'
    ROW_ALT_COLOR = 'FFEFF6FF'

    wb = Workbook()
    ws = wb.active
    ws.title = 'Заявки'

    headers = ['ID', 'Тема', 'Клиент', 'Менеджер', 'Статус', 'Приоритет', 'Услуга', 'Дата создания']
    header_fill = PatternFill(fgColor=HEADER_COLOR, fill_type='solid')
    header_font = Font(color='FFFFFFFF', bold=True)

    for col_idx, header_text in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx, value=header_text)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center')

    tickets = (
        Project.objects
        .select_related('user', 'manager', 'priority', 'service')
        .order_by('-created_at')
    )

    for row_idx, ticket in enumerate(tickets, 2):
        values = [
            ticket.pk,
            ticket.title,
            f'{ticket.user.last_name} {ticket.user.first_name}'.strip() or ticket.user.username,
            (
                f'{ticket.manager.last_name} {ticket.manager.first_name}'.strip()
                if ticket.manager else '—'
            ),
            STATUS_LABELS.get(ticket.status, ticket.status),
            ticket.priority.name if ticket.priority else '—',
            ticket.service.title if ticket.service else '—',
            ticket.created_at.strftime('%d.%m.%Y %H:%M'),
        ]
        for col_idx, val in enumerate(values, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=val)
            if row_idx % 2 == 0:
                cell.fill = PatternFill(fgColor=ROW_ALT_COLOR, fill_type='solid')

    # Ширина колонок
    col_widths = [6, 36, 22, 22, 14, 14, 22, 20]
    for col_idx, width in enumerate(col_widths, 1):
        ws.column_dimensions[
            ws.cell(row=1, column=col_idx).column_letter
        ].width = width

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer,
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    )
    response['Content-Disposition'] = 'attachment; filename="tickets_report.xlsx"'
    return response


@manager_required
def export_ticket_docx_view(request, pk):
    """
    Экспорт карточки заявки в .docx (F10).
    Содержит реквизиты заявки, описание и историю комментариев.
    """
    import io
    from docx import Document as DocxDocument
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    ticket = get_object_or_404(Project, pk=pk)
    STATUS_LABELS = dict(Project.STATUS_CHOICES)

    doc = DocxDocument()

    # Шрифт по умолчанию
    for style in doc.styles:
        try:
            style.font.name = 'Times New Roman'
        except Exception:
            pass

    # Заголовок
    heading = doc.add_heading(f'Карточка заявки № {ticket.pk}', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Таблица реквизитов
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    def add_row(label, value):
        row = table.add_row()
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = str(value)

    add_row('Тема', ticket.title)
    add_row('Клиент', ticket.user.get_full_name() or ticket.user.username)
    manager_name = (
        ticket.manager.get_full_name() or ticket.manager.username
        if ticket.manager else 'Не назначен'
    )
    add_row('Менеджер', manager_name)
    add_row('Статус', STATUS_LABELS.get(ticket.status, ticket.status))
    add_row('Приоритет', ticket.priority.name if ticket.priority else '—')
    add_row('Услуга', ticket.service.title if ticket.service else '—')
    add_row('Дата создания', ticket.created_at.strftime('%d.%m.%Y %H:%M'))
    add_row('Описание', ticket.description or '—')

    # Комментарии
    comments = ticket.comments.select_related('user').order_by('created_at')
    if comments.exists():
        doc.add_paragraph()
        doc.add_heading('История комментариев', level=2)
        for comment in comments:
            author = comment.user.get_full_name() or comment.user.username
            dt = comment.created_at.strftime('%d.%m.%Y %H:%M')
            p = doc.add_paragraph()
            p.add_run(f'{author} ({dt}): ').bold = True
            p.add_run(comment.text)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer,
        content_type=(
            'application/vnd.openxmlformats-officedocument'
            '.wordprocessingml.document'
        ),
    )
    filename = f'ticket_{ticket.pk}_report.docx'
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


# ─── Обработчики ошибок (ВКР-036) ───────────────────────────────────────────

def handler404_view(request, exception=None):
    """
    Страница 404 — ресурс не найден.
    Регистрируется в keypartner/urls.py как handler404.
    Работает только при DEBUG=False.
    """
    return render(request, '404.html', status=404)
